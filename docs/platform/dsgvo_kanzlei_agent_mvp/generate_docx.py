#!/usr/bin/env python3
"""
Generate professional DOCX version of DSGVO Sales Brief.
Brand: sevenlayers / vc83 — accent #E8520A, dark palette.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Brand tokens ──────────────────────────────────────────────
ACCENT       = RGBColor(0xE8, 0x52, 0x0A)  # #E8520A warm orange
ACCENT_DARK  = RGBColor(0xCC, 0x47, 0x09)  # #CC4709
TEXT_PRIMARY  = RGBColor(0x1A, 0x1A, 0x1A)  # near-black for print
TEXT_SECONDARY = RGBColor(0x55, 0x55, 0x55)
TEXT_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
BG_LIGHT     = RGBColor(0xF7, 0xF7, 0xF7)
BORDER_GRAY  = RGBColor(0xDD, 0xDD, 0xDD)
TABLE_HEADER_BG = "E8520A"
TABLE_ALT_BG = "FDF2EC"
TABLE_BORDER = "DDDDDD"

FONT_PRIMARY = "Calibri"  # Professional sans-serif fallback (Geist not available in DOCX)
FONT_MONO    = "Consolas"

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.abspath(os.path.join(SCRIPT_DIR, "..", "..", ".."))
LOGO_PATH = os.path.join(PROJECT_ROOT, "public", "sevenlayers-logo.png")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "SALES_BRIEF_DSGVO_DE.docx")


# ── Helpers ───────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="DDDDDD", sz="4"):
    """Set thin borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a brand-styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = TEXT_WHITE
        run.font.name = FONT_PRIMARY
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, TABLE_HEADER_BG)
        set_cell_borders(cell, TABLE_HEADER_BG, "6")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = TEXT_PRIMARY
            run.font.name = FONT_PRIMARY
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_cell_borders(cell, TABLE_BORDER)
            if r_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


def add_heading(doc, text, level=1):
    """Add a heading with brand styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_PRIMARY
        run.font.color.rgb = ACCENT if level <= 2 else TEXT_PRIMARY
    return h


def add_para(doc, text, bold=False, italic=False, size=10, color=None, align=None, space_after=Pt(6)):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = FONT_PRIMARY
    run.font.color.rgb = color or TEXT_PRIMARY
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.name = FONT_PRIMARY
        run.font.color.rgb = TEXT_PRIMARY
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = FONT_PRIMARY
    run.font.color.rgb = TEXT_PRIMARY
    p.paragraph_format.space_after = Pt(2)
    return p


def add_accent_divider(doc):
    """Add a thin accent-colored divider line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Use a border-bottom on the paragraph
    pPr = p._p.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{TABLE_HEADER_BG}"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)


def add_callout_box(doc, text):
    """Add a highlighted callout box (shaded paragraph)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = FONT_PRIMARY
    run.font.color.rgb = ACCENT_DARK
    run.italic = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Add shading
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TABLE_ALT_BG}"/>')
    pPr.append(shading)
    # Add left border accent
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="4" w:color="{TABLE_HEADER_BG}"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)
    return p


# ── Document builder ──────────────────────────────────────────
def build_document():
    doc = Document()

    # ── Page setup ────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21)    # A4
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Default font ──────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = FONT_PRIMARY
    style.font.size = Pt(10)
    style.font.color.rgb = TEXT_PRIMARY
    style.paragraph_format.space_after = Pt(6)

    # Style headings
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = FONT_PRIMARY
        hs.font.color.rgb = ACCENT if level <= 2 else TEXT_PRIMARY

    # ══════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════

    # Logo
    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Cm(5))
        p.paragraph_format.space_after = Pt(40)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("DSGVO-Konformität\nunserer Plattform")
    run.font.size = Pt(28)
    run.font.color.rgb = ACCENT
    run.font.name = FONT_PRIMARY
    run.bold = True
    p.paragraph_format.space_after = Pt(12)

    add_accent_divider(doc)

    # Subtitle
    p = doc.add_paragraph()
    run = p.add_run("Datenschutz & Berufsgeheimnisschutz für Kanzleien")
    run.font.size = Pt(14)
    run.font.color.rgb = TEXT_SECONDARY
    run.font.name = FONT_PRIMARY
    p.paragraph_format.space_after = Pt(24)

    # Description
    add_para(doc,
        "Dieses Dokument richtet sich an Kanzleien (Steuerberater, Rechtsanwälte, "
        "Wirtschaftsprüfer), die unsere Plattform mit KI-Assistenzfunktionen einsetzen "
        "möchten. Es fasst die technischen, organisatorischen und rechtlichen Maßnahmen "
        "zusammen, die den Schutz personenbezogener Daten und beruflicher Geheimnisse "
        "gewährleisten.",
        size=10, color=TEXT_SECONDARY)

    # Meta
    add_para(doc, "Vertraulich  |  Stand: März 2026  |  Version 1.0",
             size=9, color=TEXT_SECONDARY, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_after=Pt(4))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Inhaltsverzeichnis", level=1)
    toc_items = [
        "1. Unser Datenschutz-Ansatz: Fail-Closed by Design",
        "2. Schutz beruflicher Geheimnisse (§ 203 StGB / § 62a StBerG)",
        "3. Technische und organisatorische Maßnahmen (TOMs)",
        "4. Auftragsverarbeitung (AVV nach Art. 28 DSGVO)",
        "5. Betroffenenrechte (Art. 12–23 DSGVO)",
        "6. Sicherheitsvorfälle & Meldepflichten",
        "7. Datenverarbeitung im Überblick",
        "8. Verfügbare Compliance-Dokumentation",
        "9. Häufige Fragen (FAQ)",
        "10. Zusammenfassung",
    ]
    for item in toc_items:
        add_para(doc, item, size=10, color=TEXT_PRIMARY, space_after=Pt(4))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # SECTION 1 — Fail-Closed by Design
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "1. Unser Datenschutz-Ansatz: Fail-Closed by Design", level=1)

    add_para(doc,
        "Unsere Plattform verfolgt einen Fail-Closed-Ansatz: Im Zweifelsfall wird der "
        "Zugang gesperrt, nicht geöffnet. Das bedeutet konkret:")

    add_bullet(doc, " Kein Dienst wird freigeschaltet, bevor eine vollständige datenschutzrechtliche Prüfung abgeschlossen ist", bold_prefix="●")
    add_bullet(doc, " Keine Datenverarbeitung ohne dokumentierte Rechtsgrundlage (Art. 6 DSGVO)", bold_prefix="●")
    add_bullet(doc, " Kein KI-gestützter Vorgang ohne menschliche Freigabe (Human-in-the-Loop)", bold_prefix="●")
    add_bullet(doc, " Keine Weitergabe von Mandantendaten an Dritte ohne vertragliche Absicherung (AVV nach Art. 28 DSGVO)", bold_prefix="●")

    add_callout_box(doc,
        "Für Sie bedeutet das: Selbst wenn eine Funktion technisch verfügbar wäre, "
        "wird sie erst nach vollständiger Prüfung und Freigabe aktiviert.")

    # ══════════════════════════════════════════════════════════
    # SECTION 2 — Berufsgeheimnisschutz
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "2. Schutz beruflicher Geheimnisse (§ 203 StGB / § 62a StBerG)", level=1)

    add_para(doc,
        "Als Kanzlei unterliegen Sie besonderen Verschwiegenheitspflichten. "
        "Unsere Plattform berücksichtigt diese vollständig:")

    add_heading(doc, "Datenklassifizierung", level=2)

    add_styled_table(doc,
        ["Klasse", "Beschreibung", "Schutzmaßnahmen"],
        [
            ["Öffentlich", "Allgemein zugängliche Informationen", "Standardschutz"],
            ["Intern", "Betriebliche Daten ohne Mandantenbezug", "Zugriffskontrolle, Protokollierung"],
            ["Vertraulich", "Personenbezogene Daten, Geschäftsdaten", "Verschlüsselung, RBAC, Audit-Log"],
            ["Berufsgeheimnis", "Mandantenspezifische Daten", "Höchste Schutzklasse: Verschlüsselung, RBAC + MFA, menschliche Freigabe, vollständige Protokollierung"],
        ],
        col_widths=[3.5, 5, 7.5]
    )

    add_heading(doc, "Verbindliche Regeln für den KI-Assistenten", level=2)

    rules = [
        ("Keine autonome Kommunikation – ", "Der KI-Assistent versendet keine Nachrichten, E-Mails oder Dokumente ohne Ihre explizite Freigabe."),
        ("Keine autonome Rechts- oder Steuerberatung – ", "Der Assistent erstellt Entwürfe und Vorschläge; die fachliche Verantwortung bleibt beim Berufsträger."),
        ("Kein Modelltraining mit Ihren Daten – ", "Ihre Daten werden nicht zum Training von KI-Modellen verwendet."),
        ("Datenminimierung in Prompts – ", "Es werden nur die für den jeweiligen Vorgang erforderlichen Daten an das KI-Modell übermittelt."),
        ("Audit-Protokollierung – ", "Jede KI-Interaktion wird revisionssicher protokolliert."),
    ]
    for i, (prefix, text) in enumerate(rules, 1):
        add_bullet(doc, text, bold_prefix=f"{i}. {prefix}")

    # ══════════════════════════════════════════════════════════
    # SECTION 3 — TOMs
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "3. Technische und organisatorische Maßnahmen (TOMs)", level=1)

    add_para(doc,
        "Unsere Plattform implementiert umfassende Schutzmaßnahmen gemäß Art. 32 DSGVO:")

    add_heading(doc, "Zugriffskontrolle & Authentifizierung", level=2)
    add_bullet(doc, " Granulare Berechtigungen pro Benutzer und Organisation", bold_prefix="Rollenbasierte Zugriffskontrolle (RBAC) –")
    add_bullet(doc, " Zusätzliche Absicherung des Zugangs", bold_prefix="Multi-Faktor-Authentifizierung (MFA) –")
    add_bullet(doc, " Strikte Datenisolierung zwischen Organisationen", bold_prefix="Mandantentrennung –")

    add_heading(doc, "Verschlüsselung & Transport", level=2)
    add_bullet(doc, " für alle Datenübertragungen (in transit)", bold_prefix="TLS-Verschlüsselung")
    add_bullet(doc, " gespeicherter Daten (at rest)", bold_prefix="Verschlüsselung")
    add_bullet(doc, " zwischen allen Systemkomponenten", bold_prefix="Sichere API-Kommunikation")

    add_heading(doc, "Protokollierung & Nachvollziehbarkeit", level=2)
    add_bullet(doc, " – Alle relevanten Vorgänge werden automatisch protokolliert", bold_prefix="Compliance-Audit-Log")
    add_bullet(doc, " – Protokolle können nicht nachträglich manipuliert werden", bold_prefix="Unveränderliche Audit-Einträge")
    add_bullet(doc, " jeder KI-gestützten Aktion mit Zeitstempel, Benutzer und Kontext", bold_prefix="Nachvollziehbarkeit")

    add_heading(doc, "Verfügbarkeit & Wiederherstellung", level=2)
    add_bullet(doc, " der Datenbank", bold_prefix="Automatische Backups")
    add_bullet(doc, " für den Notfall", bold_prefix="Definierte Wiederherstellungsprozesse")
    add_bullet(doc, " bei Systemstörungen", bold_prefix="Überwachung und Alarmierung")

    # ══════════════════════════════════════════════════════════
    # SECTION 4 — AVV
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "4. Auftragsverarbeitung (AVV nach Art. 28 DSGVO)", level=1)

    add_para(doc,
        "Für jeden Dienstleister, der personenbezogene Daten verarbeitet, wird eine "
        "Auftragsverarbeitungsvereinbarung (AVV) abgeschlossen. Unsere AVV umfasst:")

    avv_items = [
        ("Verarbeitungszwecke und Datenkategorien – ", "Exakte Festlegung, welche Daten wofür verarbeitet werden"),
        ("Vertraulichkeitsverpflichtung – ", "Erweitert um die Anforderungen des § 203 StGB und § 62a StBerG"),
        ("Unterauftragnehmer-Management – ", "Transparente Liste aller Unterauftragnehmer mit Widerspruchsrecht bei Änderungen"),
        ("Drittlandtransfers – ", "Dokumentierte Schutzmaßnahmen gemäß Art. 44 ff. DSGVO (SCCs, Angemessenheitsbeschlüsse)"),
        ("Unterstützung bei Betroffenenrechten – ", "Technische Unterstützung für Auskunft, Löschung, Berichtigung etc."),
        ("Meldepflichten – ", "Benachrichtigung bei Sicherheitsvorfällen innerhalb von 24 Stunden"),
    ]
    for prefix, text in avv_items:
        add_bullet(doc, text, bold_prefix=prefix)

    # ══════════════════════════════════════════════════════════
    # SECTION 5 — Betroffenenrechte
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "5. Betroffenenrechte (Art. 12–23 DSGVO)", level=1)

    add_para(doc,
        "Wir unterstützen Sie bei der Erfüllung Ihrer Pflichten gegenüber betroffenen Personen:")

    add_styled_table(doc,
        ["Recht", "Umsetzung"],
        [
            ["Auskunft (Art. 15)", "Export aller gespeicherten Daten einer betroffenen Person"],
            ["Berichtigung (Art. 16)", "Korrektur unrichtiger Daten über die Plattform"],
            ["Löschung (Art. 17)", "Löschung personenbezogener Daten auf Anfrage"],
            ["Einschränkung (Art. 18)", "Sperrung der Verarbeitung bei laufender Prüfung"],
            ["Datenübertragbarkeit (Art. 20)", "Bereitstellung der Daten in maschinenlesbarem Format"],
            ["Widerspruch (Art. 21)", "Einstellung der Verarbeitung bei berechtigtem Widerspruch"],
        ],
        col_widths=[5, 11]
    )

    add_heading(doc, "Bearbeitungszeiten", level=2)
    add_bullet(doc, " ≤ 2 Werktage", bold_prefix="Eingangsbestätigung:")
    add_bullet(doc, " ≤ 5 Werktage", bold_prefix="Identitätsprüfung:")
    add_bullet(doc, " ≤ 30 Kalendertage", bold_prefix="Vollständige Bearbeitung:")

    # ══════════════════════════════════════════════════════════
    # SECTION 6 — Incident Response
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "6. Sicherheitsvorfälle & Meldepflichten", level=1)

    add_para(doc,
        "Unser Incident-Response-Prozess ist auf die besonderen Anforderungen "
        "von Kanzleien ausgerichtet:")

    add_heading(doc, "Reaktionszeiten nach Schweregrad", level=2)

    add_styled_table(doc,
        ["Schweregrad", "Reaktionszeit", "Beispiel"],
        [
            ["SEV-1 (Kritisch)", "≤ 15 Minuten", "Verdacht auf Zugriff auf Mandantendaten / Berufsgeheimnisse"],
            ["SEV-2 (Hoch)", "≤ 30 Minuten", "Sicherheitsvorfall mit möglichem Personenbezug"],
            ["SEV-3 (Mittel)", "≤ 4 Stunden", "Sicherheitsvorfall ohne Personenbezug"],
        ],
        col_widths=[4, 3.5, 8.5]
    )

    add_heading(doc, "Meldepflichten gemäß Art. 33/34 DSGVO", level=2)
    add_bullet(doc, " an die Aufsichtsbehörde innerhalb von 72 Stunden nach Kenntnisnahme", bold_prefix="Meldung")
    add_bullet(doc, " der Betroffenen bei hohem Risiko", bold_prefix="Benachrichtigung")
    add_bullet(doc, " jedes Vorfalls mit lückenloser Nachverfolgung", bold_prefix="Dokumentation")
    add_bullet(doc, " mit konkreten Verbesserungsmaßnahmen", bold_prefix="Postmortem-Analyse")

    # ══════════════════════════════════════════════════════════
    # SECTION 7 — Datenverarbeitung
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "7. Datenverarbeitung im Überblick", level=1)

    add_heading(doc, "Verarbeitungszwecke und Rechtsgrundlagen", level=2)

    add_styled_table(doc,
        ["Zweck", "Rechtsgrundlage", "Datenkategorien"],
        [
            ["Kontoverwaltung", "Art. 6 (1)(b) – Vertragsdurchführung", "Name, E-Mail, Kontaktdaten"],
            ["KI-Assistenz", "Art. 6 (1)(b) – Vertragsdurchführung", "Eingabedaten, Entwürfe (minimiert)"],
            ["Abrechnung", "Art. 6 (1)(b) + (c) – Vertrag + gesetzl. Pflicht", "Zahlungsdaten, Rechnungsadresse"],
            ["Sicherheit", "Art. 6 (1)(f) – Berechtigtes Interesse", "Zugriffsprotokolle, IP-Adressen"],
        ],
        col_widths=[3.5, 6.5, 6]
    )

    add_heading(doc, "Ihre Daten werden NICHT verwendet für:", level=2)
    add_bullet(doc, " Training von KI-Modellen", bold_prefix="✗")
    add_bullet(doc, " Profilbildung oder automatisierte Entscheidungsfindung", bold_prefix="✗")
    add_bullet(doc, " Weitergabe an Dritte zu Werbezwecken", bold_prefix="✗")
    add_bullet(doc, " Zwecke außerhalb des Vertragsverhältnisses", bold_prefix="✗")

    # ══════════════════════════════════════════════════════════
    # SECTION 8 — Compliance-Dokumentation
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "8. Verfügbare Compliance-Dokumentation", level=1)

    add_para(doc, "Folgende Dokumente stellen wir Ihnen auf Anfrage zur Verfügung:")

    add_styled_table(doc,
        ["Dokument", "Inhalt"],
        [
            ["Datenschutzerklärung", "Vollständige Information nach Art. 13/14 DSGVO"],
            ["AVV (Auftragsverarbeitungsvereinbarung)", "Vertrag nach Art. 28 DSGVO inkl. § 203-Klauseln"],
            ["TOM-Verzeichnis", "Technische und organisatorische Maßnahmen nach Art. 32 DSGVO"],
            ["Unterauftragnehmer-Liste", "Aktuelle Liste aller Unterauftragnehmer mit Verarbeitungszwecken"],
            ["Drittlandtransfer-Register", "Dokumentation aller Datentransfers außerhalb des EWR"],
            ["KI-Nutzungsrichtlinie", "Verbindliche Regeln für den Einsatz des KI-Assistenten"],
        ],
        col_widths=[6, 10]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION 9 — FAQ
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "9. Häufige Fragen (FAQ)", level=1)

    faqs = [
        ("Werden meine Mandantendaten zum KI-Training verwendet?",
         "Nein. Wir schließen die Verwendung Ihrer Daten zum Training von KI-Modellen vertraglich aus. "
         "Dies ist in unserer AVV und der KI-Nutzungsrichtlinie verbindlich festgelegt."),
        ("Kann der KI-Assistent eigenständig Mandanten kontaktieren?",
         "Nein. Der KI-Assistent arbeitet ausschließlich als Entwurfs- und Strukturierungshilfe. "
         "Jede externe Kommunikation erfordert Ihre explizite Freigabe als Berufsträger."),
        ("Wo werden meine Daten gespeichert?",
         "Alle Datenverarbeitungen sind in unserem Unterauftragnehmer-Verzeichnis dokumentiert. "
         "Für Drittlandtransfers sind geeignete Garantien gemäß Art. 44 ff. DSGVO implementiert "
         "(z. B. Standardvertragsklauseln)."),
        ("Wie erfülle ich meine Auskunftspflichten gegenüber Mandanten?",
         "Unsere Plattform unterstützt Sie mit Exportfunktionen und definierten Prozessen für alle "
         "Betroffenenrechte nach Art. 12–23 DSGVO. Unser Betriebsteam unterstützt bei der technischen Umsetzung."),
        ("Was passiert bei einem Sicherheitsvorfall?",
         "Unser Incident-Response-Prozess sieht eine Reaktionszeit von ≤ 15 Minuten bei kritischen "
         "Vorfällen vor. Sie werden innerhalb von 24 Stunden informiert. Die Meldung an die "
         "Aufsichtsbehörde erfolgt fristgerecht innerhalb von 72 Stunden."),
        ("Erfüllt die Plattform die Anforderungen des § 62a StBerG?",
         "Ja. Unsere AVV enthält spezifische Klauseln zum Schutz beruflicher Geheimnisse gemäß "
         "§ 203 StGB und § 62a StBerG. Alle Unterauftragnehmer werden auf diese Anforderungen "
         "geprüft und verpflichtet."),
    ]
    for q, a in faqs:
        add_para(doc, q, bold=True, size=10, color=ACCENT_DARK, space_after=Pt(2))
        add_para(doc, a, size=9.5, color=TEXT_PRIMARY, space_after=Pt(12))

    # ══════════════════════════════════════════════════════════
    # SECTION 10 — Zusammenfassung
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "10. Zusammenfassung", level=1)

    add_para(doc, "Warum unsere Plattform für Kanzleien geeignet ist:", bold=True, size=10)

    add_styled_table(doc,
        ["Anforderung", "Umsetzung"],
        [
            ["DSGVO-Konformität", "Vollständiges Compliance-Framework nach Art. 5, 6, 12–23, 28, 32, 33/34 DSGVO"],
            ["Berufsgeheimnisschutz", "Spezifische Maßnahmen für § 203 StGB und § 62a StBerG"],
            ["Human-in-the-Loop", "Keine autonome Außenkommunikation oder Beratung"],
            ["Kein Modelltraining", "Vertraglicher Ausschluss der Datennutzung zum Training"],
            ["Datenminimierung", "Nur erforderliche Daten werden an KI-Modelle übermittelt"],
            ["Audit-Fähigkeit", "Lückenlose, revisionssichere Protokollierung"],
            ["Mandantentrennung", "Strikte Datenisolierung zwischen Organisationen"],
            ["Incident Response", "Definierte Prozesse mit SLAs (≤ 15 Min. bei SEV-1)"],
            ["Betroffenenrechte", "Vollständige Unterstützung aller DSGVO-Rechte"],
            ["Transparenz", "Offene Dokumentation aller Verarbeitungen und Unterauftragnehmer"],
        ],
        col_widths=[5, 11]
    )

    # ══════════════════════════════════════════════════════════
    # CONTACT / FOOTER
    # ══════════════════════════════════════════════════════════
    add_accent_divider(doc)

    add_heading(doc, "Kontakt", level=2)

    contact_lines = [
        ("Datenschutz: ", "[E-Mail-Adresse einfügen]"),
        ("Vertrieb: ", "[E-Mail-Adresse einfügen]"),
        ("Technischer Support: ", "[E-Mail-Adresse einfügen]"),
    ]
    for prefix, text in contact_lines:
        add_bullet(doc, text, bold_prefix=prefix)

    add_accent_divider(doc)

    add_para(doc,
        "Dieses Dokument dient der Information und ersetzt keine individuelle Rechtsberatung. "
        "Stand: März 2026.",
        size=8, color=TEXT_SECONDARY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Save ──────────────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    print(f"✓ Generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
