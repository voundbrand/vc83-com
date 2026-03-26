#!/usr/bin/env python3
"""
Generate clean .docx files from Phase 0 markdown sources.
No dark backgrounds, no full-page formatting — clean white pages with professional typography.
"""

import re
import os
import sys

# Ensure python-docx is importable
sys.path.insert(0, "/Users/foundbrand_001/Library/Python/3.9/lib/python/site-packages")

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_DIR = os.path.join(SCRIPT_DIR, "docx")

MD_FILES = [
    "00_PHASE_0_OVERVIEW.md",
    "01_PHASE_0_ICP.md",
    "02_PHASE_0_SALES_MOTION.md",
    "03_PHASE_0_PRICING.md",
    "04_PHASE_0_DEMO.md",
    "05_PHASE_0_BRIDGE.md",
]


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_document(doc):
    """Set up clean, professional document styles."""
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Style: Normal
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading styles
    for level, (size, color_hex) in enumerate(
        [
            (Pt(22), "1a1a2e"),  # H1
            (Pt(16), "2d2d44"),  # H2
            (Pt(13), "3d3d5c"),  # H3
        ],
        start=1,
    ):
        hstyle = doc.styles[f"Heading {level}"]
        hstyle.font.name = "Calibri"
        hstyle.font.size = size
        hstyle.font.bold = True
        hstyle.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16),
        )
        hstyle.paragraph_format.space_before = Pt(18 if level == 1 else 14)
        hstyle.paragraph_format.space_after = Pt(8)


def add_table(doc, header_row, data_rows):
    """Add a clean table with light header shading."""
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    for i, text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text.strip())
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, "2d2d44")

    # Data rows
    for r_idx, row_data in enumerate(data_rows):
        for c_idx, text in enumerate(row_data):
            if c_idx >= cols:
                break
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            # Handle bold markers
            parts = re.split(r"\*\*(.+?)\*\*", text.strip())
            for j, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.size = Pt(9.5)
                run.font.name = "Calibri"
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                if j % 2 == 1:
                    run.bold = True
            # Alternate row shading
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F5F5F8")

    # Column widths - distribute evenly
    doc.add_paragraph()  # spacing after table


def parse_table_block(lines):
    """Parse markdown table lines into header + data rows."""
    # Filter out separator lines (|---|---|)
    content_lines = [l for l in lines if not re.match(r"^\|[\s\-:|]+\|$", l)]
    if not content_lines:
        return None, None

    def split_row(line):
        cells = line.strip().strip("|").split("|")
        return [c.strip() for c in cells]

    header = split_row(content_lines[0])
    data = [split_row(l) for l in content_lines[1:]]
    return header, data


def convert_md_to_docx(md_path, docx_path):
    """Convert a markdown file to a clean docx."""
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    style_document(doc)

    lines = content.split("\n")
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block - add as formatted paragraph
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                # Light background via shading
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="F0F0F4" w:val="clear"/>'
                )
                p._p.get_or_add_pPr().append(shading)
                in_code_block = False
                code_lines = []
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table detection
        if line.strip().startswith("|") and i + 1 < len(lines):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            header, data = parse_table_block(table_lines)
            if header and data is not None:
                add_table(doc, header, data)
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # Horizontal rule → just add spacing, no page break
        if line.strip() == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Thin line
            run = p.add_run("─" * 60)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Blockquote
        if line.strip().startswith(">"):
            quote_lines = []
            while i < len(lines) and (
                lines[i].strip().startswith(">") or lines[i].strip() == ""
            ):
                text = lines[i].strip()
                if text.startswith("> "):
                    quote_lines.append(text[2:])
                elif text == ">":
                    quote_lines.append("")
                elif text == "" and quote_lines:
                    # Check if next line is still a quote
                    if i + 1 < len(lines) and lines[i + 1].strip().startswith(">"):
                        quote_lines.append("")
                    else:
                        break
                else:
                    break
                i += 1

            quote_text = "\n".join(quote_lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

            # Parse bold within blockquote
            parts = re.split(r"\*\*(.+?)\*\*", quote_text)
            for j, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                run.italic = True
                if j % 2 == 1:
                    run.bold = True
            continue

        # Bullet list
        if line.strip().startswith("- "):
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            # Handle bold
            parts = re.split(r"\*\*(.+?)\*\*", text)
            p.clear()
            for j, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.name = "Calibri"
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                if j % 2 == 1:
                    run.bold = True
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\d+)\.\s+(.+)$", line.strip())
        if m:
            text = m.group(2)
            p = doc.add_paragraph(style="List Number")
            p.clear()
            parts = re.split(r"\*\*(.+?)\*\*", text)
            for j, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.name = "Calibri"
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                if j % 2 == 1:
                    run.bold = True
            i += 1
            continue

        # Italic footer lines
        if line.strip().startswith("*") and line.strip().endswith("*") and not line.strip().startswith("**"):
            text = line.strip().strip("*")
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            i += 1
            continue

        # Regular paragraph
        text = line.strip()
        if text:
            p = doc.add_paragraph()
            parts = re.split(r"\*\*(.+?)\*\*", text)
            for j, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.name = "Calibri"
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                if j % 2 == 1:
                    run.bold = True

        i += 1

    doc.save(docx_path)
    print(f"  ✓ {os.path.basename(docx_path)}")


def main():
    os.makedirs(DOCX_DIR, exist_ok=True)
    print("Generating clean docx files...\n")

    for md_file in MD_FILES:
        md_path = os.path.join(SCRIPT_DIR, md_file)
        docx_name = md_file.replace(".md", ".docx")
        docx_path = os.path.join(DOCX_DIR, docx_name)

        if not os.path.exists(md_path):
            print(f"  ✗ {md_file} not found, skipping")
            continue

        convert_md_to_docx(md_path, docx_path)

    print(f"\nDone! {len(MD_FILES)} files generated in {DOCX_DIR}")


if __name__ == "__main__":
    main()
